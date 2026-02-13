"""PDF and DOCX document loaders."""

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from loguru import logger
from pypdf import PdfReader


@dataclass
class DocumentPage:
    """Represents a single page/section of a document."""

    text: str
    page_number: int
    metadata: dict = field(default_factory=dict)


@dataclass
class LoadedDocument:
    """Full document with all pages and metadata."""

    pages: list[DocumentPage]
    filename: str
    file_type: str
    total_pages: int

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


def load_pdf(file_path: str | Path) -> LoadedDocument:
    """Load a PDF file and extract text page by page."""
    path = Path(file_path)
    logger.info(f"Loading PDF: {path.name}")

    reader = PdfReader(str(path))
    pages: list[DocumentPage] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(
                DocumentPage(
                    text=text,
                    page_number=i + 1,
                    metadata={"source": path.name, "page": i + 1},
                )
            )

    logger.info(f"Loaded {len(pages)} pages from {path.name}")
    return LoadedDocument(
        pages=pages,
        filename=path.name,
        file_type="pdf",
        total_pages=len(reader.pages),
    )


def load_docx(file_path: str | Path) -> LoadedDocument:
    """Load a DOCX file and extract text from paragraphs, tables, headers, and footers."""
    path = Path(file_path)
    logger.info(f"Loading DOCX: {path.name}")

    doc = DocxDocument(str(path))

    # Extract paragraphs (standard API handles all the XML complexity)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract table contents (paragraph-by-paragraph from each cell)
    # Many legal DOCX files store all content inside a single table,
    # so we extract individual paragraphs from cells for proper chunking.
    seen_cell_ids: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # python-docx can return the same cell object for merged cells
                if id(cell) in seen_cell_ids:
                    continue
                seen_cell_ids.add(id(cell))
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

    # Extract headers and footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer and header_footer.paragraphs:
                for para in header_footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

    logger.info(f"Extracted {len(paragraphs)} text blocks from {path.name}")

    if not paragraphs:
        logger.warning(f"No text extracted from {path.name}")


    # Group paragraphs into logical pages (~3000 chars each)
    pages: list[DocumentPage] = []
    current_text = ""
    page_num = 1

    for para in paragraphs:
        current_text += para + "\n\n"
        if len(current_text) >= 3000:
            pages.append(
                DocumentPage(
                    text=current_text.strip(),
                    page_number=page_num,
                    metadata={"source": path.name, "page": page_num},
                )
            )
            current_text = ""
            page_num += 1

    if current_text.strip():
        pages.append(
            DocumentPage(
                text=current_text.strip(),
                page_number=page_num,
                metadata={"source": path.name, "page": page_num},
            )
        )

    logger.info(f"Loaded {len(pages)} sections from {path.name}")
    return LoadedDocument(
        pages=pages,
        filename=path.name,
        file_type="docx",
        total_pages=len(pages),
    )


def load_document(file_path: str | Path) -> LoadedDocument:
    """Load a document based on its file extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf or .docx")
